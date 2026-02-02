from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.table import _Row

def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    # Junta runs para evitar placeholder quebrado em runs diferentes
    full = "".join(run.text for run in paragraph.runs)
    if not full:
        return

    changed = False
    for key, val in mapping.items():
        token = "{{" + key + "}}"
        if token in full:
            full = full.replace(token, val)
            changed = True

    if changed:
        # limpa runs e coloca tudo no primeiro run
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = full
        else:
            paragraph.add_run(full)

def _row_has_token(row, token: str) -> bool:
    for cell in row.cells:
        for p in cell.paragraphs:
            if token in "".join(run.text for run in p.runs):
                return True
    return False

def _replace_in_row(row, mapping: Dict[str, str]) -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            _replace_in_paragraph(p, mapping)

def _expand_trecho_rows(doc: Document, rows: Dict[str, list]) -> None:
    ida_rows = rows.get("ida") or []
    ret_rows = rows.get("retorno") or []
    for table in doc.tables:
        i = 0
        while i < len(table.rows):
            row = table.rows[i]
            if _row_has_token(row, "{{ida_origem}}"):
                row_data = table.rows[i + 1] if (i + 1) < len(table.rows) else None
                has_data_row = row_data is not None and _row_has_token(row_data, "{{ida_data_hora}}")
                tmpl_a = deepcopy(row._tr)
                tmpl_b = deepcopy(row_data._tr) if has_data_row else None
                if ida_rows:
                    _replace_in_row(row, ida_rows[0])
                    if has_data_row:
                        _replace_in_row(row_data, ida_rows[0])
                    prev_tr = row_data._tr if has_data_row else row._tr
                    for item in ida_rows[1:]:
                        new_a = deepcopy(tmpl_a)
                        prev_tr.addnext(new_a)
                        _replace_in_row(_Row(new_a, table), item)
                        prev_tr = new_a
                        if tmpl_b is not None:
                            new_b = deepcopy(tmpl_b)
                            prev_tr.addnext(new_b)
                            _replace_in_row(_Row(new_b, table), item)
                            prev_tr = new_b
                else:
                    _replace_in_row(row, {"ida_origem": "", "ida_destino": "", "ida_data_hora": ""})
                    if has_data_row:
                        _replace_in_row(row_data, {"ida_origem": "", "ida_destino": "", "ida_data_hora": ""})
                i += 2 if has_data_row else 1
                continue
            if _row_has_token(row, "{{retorno_origem}}"):
                row_data = table.rows[i + 1] if (i + 1) < len(table.rows) else None
                has_data_row = row_data is not None and _row_has_token(row_data, "{{retorno_data_hora}}")
                tmpl_a = deepcopy(row._tr)
                tmpl_b = deepcopy(row_data._tr) if has_data_row else None
                if ret_rows:
                    _replace_in_row(row, ret_rows[0])
                    if has_data_row:
                        _replace_in_row(row_data, ret_rows[0])
                    prev_tr = row_data._tr if has_data_row else row._tr
                    for item in ret_rows[1:]:
                        new_a = deepcopy(tmpl_a)
                        prev_tr.addnext(new_a)
                        _replace_in_row(_Row(new_a, table), item)
                        prev_tr = new_a
                        if tmpl_b is not None:
                            new_b = deepcopy(tmpl_b)
                            prev_tr.addnext(new_b)
                            _replace_in_row(_Row(new_b, table), item)
                            prev_tr = new_b
                else:
                    _replace_in_row(row, {"retorno_origem": "", "retorno_destino": "", "retorno_data_hora": ""})
                    if has_data_row:
                        _replace_in_row(row_data, {"retorno_origem": "", "retorno_destino": "", "retorno_data_hora": ""})
                i += 2 if has_data_row else 1
                continue
            i += 1


def render_docx_from_template(
    template_path: Path,
    output_path: Path,
    mapping: Dict[str, str],
    rows: Optional[Dict[str, list]] = None,
) -> None:
    doc = Document(str(template_path))

    if rows:
        _expand_trecho_rows(doc, rows)

    # parágrafos
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    # tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
